from docx import Document


class Paragraph_cords:
	def __init__(self, p, c):
		self.p = p
		self.c = c

class DocManager:

	@staticmethod
	def assemble_from_flags(doc_name: str, flags: list, parent_docuemnt_obj):
		pass

	def __init__(self, document_name: str):
		self.name = document_name
		self.docx = Document(docx=document_name)
		self.paragraphs = self.docx.paragraphs
		self.flags = []
		
	def flag_markers(self, kwargs) -> dict: 

		for (arg, run_cond) in kwargs:

			if run_cond is True:
				
				if arg == "d":
					temp = {"mark": "--",
							"color": "orange",
							"found_at": []}
				
					self.flags.append(temp)
				elif arg == "i":
					temp = {"mark": "\"",
							"color": "yellow",
							"found_at": []}
					self.flags.append(temp)

				elif arg == "q":
					temp = {"mark": "\'",
							"color": "grey",
							"found_at": []}
					self.flags.append(temp)

				else: # is s
					
					msg = "What words to search for (\"Exit\" to exit): "
					temp = input("msg")
					search_words = []

					while temp != "Exit":
						
						search_words.append(temp)
						temp = input("msg")

					if search_words != []:
						print(f"Will search for words:")

						for word in search_words:
							print(f"    -{word}")

							temp = {"mark": word,
									"color": "blue",
									"found_at": []}

							self.flags.append(temp)

	def temp_function(self):

		file_str = ""
		delimiter = "^"

		for paragraph in self.paragraphs.text:
			file_str += f"{paragraph}{delimiter}"

		sentences = file_str.split(".")

		"""
		global_str = ""
		delimiter = "^"

		search_for = "test"

		for paragraph in paragraphs.text:
		global_str += f"{paragraph}{delimiter}"
		
		sentences = global_str.split(".")
		flags = {"quotes": [],
				"internal": [],
				"dashed": [],
				"searchWord": []}

		count = 0

		for sentence in sentences:

		if "\"" in sentence:
			flags["quotes"].append(count)
		
		elif "\'" in sentence:
			flags["internal"].append(count)
			
		elif "--" in sentnece:
			flags["dashed"].append(count)
			
		if search_for in sentence:
			flags["searchWord"].append(count)
			
		count += 1
		"""

	def save(self) -> None:
		pass
